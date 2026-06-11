from docx import Document #This imports a library that knows how to read Microsoft Word files
from pathlib import Path
# Input and output paths
input_file = Path("data/raw/oran/O-RAN.WG1.TS.OAD-R005-v16.00.docx")
output_file = Path("data/processed/oran.txt")
clean_output = Path("data/processed/oran_clean.txt")

# loads the entire 3GPP specification into memory.
doc = Document(input_file)

# Extract text
text = []

for para in doc.paragraphs:
    text.append(para.text)

full_text = "\n".join(text)

clean_lines = []

for line in full_text.splitlines():
    line = " ".join(line.split())

    if line:
        clean_lines.append(line)

clean_text = "\n".join(clean_lines)

# Save extracted text
with open(output_file, "w", encoding="utf-8") as f:
    f.write(full_text)

# Save cleaned text
with open(clean_output, "w", encoding="utf-8") as f:
    f.write(clean_text)

word_count = len(full_text.split())
char_count = len(full_text)
paragraph_count = len(doc.paragraphs)

print("Extraction completed!")
print(f"Saved to: {output_file}")
print(f"Clean file saved to: {clean_output}")

print(f"Words: {word_count}")
print(f"Characters: {char_count}")
print(f"Paragraphs: {paragraph_count}")

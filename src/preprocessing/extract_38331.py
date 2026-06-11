from docx import Document

# Input DOCX file
doc_path = "data/raw/3GPP/38331-i90/38331-i90.docx"

# Open document
doc = Document(doc_path)

# Extract text
text = ""

for para in doc.paragraphs:
    text += para.text + "\n"

# Save extracted text
output_path = "data/processed/38331.txt"

with open(output_path, "w", encoding="utf-8") as f:
    f.write(text)

print("Text extracted successfully!")
print(f"Saved to: {output_path}")

word_count = len(text.split())
char_count = len(text)
paragraph_count = len(doc.paragraphs)

print("\nStatistics")
print("----------------")
print("Words:", word_count)
print("Characters:", char_count)
print("Paragraphs:", paragraph_count)

# Create cleaned text

clean_lines = []

for line in text.splitlines():
    line = " ".join(line.split())

    if line:
        clean_lines.append(line)

clean_text = "\n".join(clean_lines)

with open(
    "data/processed/38331_clean.txt",
    "w",
    encoding="utf-8"
) as f:
    f.write(clean_text)

print("Clean text file created!")